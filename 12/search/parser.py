import os
from typing import Optional


class DocumentParser:
    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.doc', '.docx'}

    def __init__(self):
        pass

    def parse(self, file_path: str) -> Optional[str]:
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}")

        if ext == '.txt':
            return self._parse_txt(file_path)
        elif ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext in ('.doc', '.docx'):
            return self._parse_docx(file_path)
        return None

    def _parse_txt(self, file_path: str) -> str:
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _parse_pdf(self, file_path: str) -> str:
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return '\n'.join(text_parts)
        except ImportError:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return '\n'.join(text_parts)
            except Exception as e:
                raise RuntimeError(f"PDF 解析失败: {e}")

    def _parse_docx(self, file_path: str) -> str:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            return '\n'.join(text_parts)
        except Exception as e:
            raise RuntimeError(f"Word 文档解析失败: {e}")

    def get_file_type(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        type_map = {
            '.txt': 'txt',
            '.pdf': 'pdf',
            '.doc': 'doc',
            '.docx': 'docx',
        }
        return type_map.get(ext, 'unknown')
